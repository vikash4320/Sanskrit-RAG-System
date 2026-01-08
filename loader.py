from langchain_community.document_loaders import TextLoader, PyPDFLoader
from docx import Document
from langchain_core.documents import Document as LCDocument

def load_documents(path):
    documents = []

    if path.endswith(".pdf"):
        loader = PyPDFLoader(path)
        documents = loader.load()

    elif path.endswith(".txt"):
        loader = TextLoader(path, encoding="utf-8")
        documents = loader.load()

    elif path.endswith(".docx"):
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        documents = [LCDocument(page_content=text)]

    else:
        raise ValueError("Unsupported file format")

    return documents
